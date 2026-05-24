"""通用工具函数模块，提供任务 ID 生成、文件操作和文档转换等功能。"""

import os
import datetime
import hashlib
import tomllib
from pathlib import Path
from app.schemas.enums import CompTemplate
from app.utils.log_util import logger
import re
import pypandoc  # type: ignore[import-unresolved]
from app.config.setting import settings

TASK_ID_PATTERN = re.compile(r"^[A-Za-z0-9][A-Za-z0-9._-]{0,127}$")


def create_task_id() -> str:
    """生成基于时间戳和随机哈希的唯一任务 ID。"""
    # 生成时间戳和随机hash
    timestamp = datetime.datetime.now().strftime("%Y%m%d-%H%M%S")
    random_hash = hashlib.md5(str(datetime.datetime.now()).encode()).hexdigest()[:8]
    return f"{timestamp}-{random_hash}"


def ensure_safe_task_id(task_id: str) -> str:
    """验证任务 ID 的合法性，防止路径遍历攻击。

    Args:
        task_id: 待验证的任务 ID。

    Returns:
        验证通过的任务 ID。

    Raises:
        ValueError: 任务 ID 不合法时抛出。
    """
    normalized = (task_id or "").strip()
    if not normalized or not TASK_ID_PATTERN.fullmatch(normalized):
        raise ValueError("非法 task_id")
    return normalized


def create_work_dir(task_id: str) -> str:
    """为指定任务创建工作目录。

    Args:
        task_id: 任务 ID。

    Returns:
        工作目录路径。
    """
    # 设置主工作目录和子目录
    work_dir = os.path.join("project", "work_dir", task_id)

    try:
        # 创建目录，如果目录已存在也不会报错
        os.makedirs(work_dir, exist_ok=True)
        return work_dir
    except Exception as e:
        # 捕获并记录创建目录时的异常
        logger.error(f"创建工作目录失败: {str(e)}")
        raise


def get_work_dir(task_id: str) -> str:
    """获取指定任务的工作目录路径。

    Args:
        task_id: 任务 ID。

    Returns:
        工作目录路径。

    Raises:
        FileNotFoundError: 工作目录不存在时抛出。
    """
    work_dir = os.path.join("project", "work_dir", task_id)
    if os.path.exists(work_dir):
        return work_dir
    else:
        logger.error(f"工作目录不存在: {work_dir}")
        raise FileNotFoundError(f"工作目录不存在: {work_dir}")


# TODO: 是不是应该将 Prompt 写成一个 class
def get_config_template(comp_template: CompTemplate = CompTemplate.CHINA) -> dict:
    """获取论文模板配置。

    Args:
        comp_template: 竞赛模板类型。

    Returns:
        模板配置字典。
    """
    if comp_template == CompTemplate.CHINA:
        return load_toml(os.path.join("app", "config", "md_template.toml"))
    return {}


def load_toml(path: str) -> dict:
    """加载 TOML 配置文件。

    Args:
        path: TOML 文件路径。
    """
    with open(path, "rb") as f:
        return tomllib.load(f)


def load_markdown(path: str) -> str:
    """加载 Markdown 文件内容。

    Args:
        path: Markdown 文件路径。
    """
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def get_current_files(folder_path: str, type: str = "all") -> list[str]:
    """获取指定目录下的文件列表。

    Args:
        folder_path: 目录路径。
        type: 文件类型过滤（all/md/ipynb/data/image）。
    """
    files = os.listdir(folder_path)
    if type == "all":
        return files
    elif type == "md":
        return [file for file in files if file.endswith(".md")]
    elif type == "ipynb":
        return [file for file in files if file.endswith(".ipynb")]
    elif type == "data":
        return [
            file for file in files if file.endswith(".xlsx") or file.endswith(".csv")
        ]
    elif type == "image":
        return [
            file for file in files if file.endswith(".png") or file.endswith(".jpg")
        ]
    return []


def transform_link(task_id: str, content: str):
    """将 Markdown 中的图片链接转换为静态资源 URL。

    Args:
        task_id: 任务 ID，用于构建 URL 路径。
        content: 包含图片链接的 Markdown 文本。
    """
    content = re.sub(
        r"!\[(.*?)\]\((.*?\.(?:png|jpg|jpeg|gif|bmp|webp))\)",
        lambda match: f"![{match.group(1)}]({settings.SERVER_HOST}/static/{task_id}/{match.group(2)})",
        content,
    )
    return content


def md_2_docx(task_id: str):
    """将 Markdown 论文转换为专业排版 DOCX。

    使用论文模板作为 reference-doc，pandoc 自动映射样式：
    - 标题、正文、页边距、字体从模板继承
    - LaTeX 公式自动转为 MathML/OMML
    - 图片通过 --resource-path 自动嵌入

    Args:
        task_id: 任务 ID。
    """
    work_dir = get_work_dir(task_id)
    md_path = os.path.join(work_dir, "res.md")
    docx_path = os.path.join(work_dir, "res.docx")

    # 论文模板路径
    template_dir = Path(__file__).parent.parent.parent / "templates"
    reference_doc = template_dir / "论文模板.docx"

    extra_args = [
        "--resource-path",
        str(work_dir),
        "--mathml",
        "--standalone",
    ]

    # 如果模板存在，使用 reference-doc 获得专业排版
    if reference_doc.exists():
        extra_args.extend(["--reference-doc", str(reference_doc)])
        logger.info(f"使用论文模板: {reference_doc}")

    # Step 1: 从 res.md 提取 HTML 表格，替换为占位符
    tables_data = _extract_html_tables(md_path)
    # 如果 res.md 已被清洗（无 HTML），从 res.json 恢复表格数据
    if not tables_data:
        tables_data = _extract_tables_from_json(work_dir)

    # Step 2: Pandoc 转换（公式→MathML，图片嵌入，正文排版）
    pypandoc.convert_file(
        source_file=md_path,
        to="docx",
        outputfile=docx_path,
        format="markdown+tex_math_dollars",
        extra_args=extra_args,
    )
    logger.info(f"DOCX 转换完成: {docx_path}")

    # Step 3: 格式化 + 插入 Word 真表格
    if reference_doc.exists():
        _sync_template_styles(str(reference_doc), docx_path, tables_data)


def md_2_pdf(task_id: str, work_dir: str | None = None) -> str | None:
    """将 res.md 转为 LaTeX 并编译为 PDF。

    集成 Markdown→LaTeX 转换器 + xelatex 编译。
    自动处理 MiKTeX 字体缓存问题。

    Args:
        task_id: 任务 ID。
        work_dir: 工作目录（可选，默认从 task_id 推算）。

    Returns:
        PDF 文件路径，编译失败返回 None。
    """
    import subprocess
    import shutil

    wd = work_dir or get_work_dir(task_id)
    md_path = os.path.join(wd, "res.md")
    tex_path = os.path.join(wd, "res.tex")
    pdf_path = os.path.join(wd, "res.pdf")

    if not os.path.exists(md_path):
        logger.warning(f"Markdown 源文件不存在: {md_path}")
        return None

    # Step 1: Markdown → LaTeX（从 res.json 恢复表格数据）
    try:
        import json, re
        from app.utils.md_to_latex import convert as md_to_latex_convert

        # 从 res.json 提取 HTML 表格（res.md 可能已被 DOCX 流程覆盖）
        import app.utils.md_to_latex as md2tex
        json_path = os.path.join(wd, "res.json")
        if os.path.exists(json_path):
            with open(json_path, "r", encoding="utf-8") as f:
                res_data = json.load(f)
            all_tables = []
            for key, val in res_data.items():
                content = val.get("response_content", "")
                tables = re.findall(r"<table>.*?</table>", content, re.DOTALL)
                all_tables.extend(tables)
            if all_tables:
                md2tex._external_tables = all_tables
                logger.info(f"从 res.json 恢复 {len(all_tables)} 个表格")

        md_to_latex_convert(md_path, tex_path, "")
        logger.info(f"LaTeX 转换完成: {tex_path}")
    except Exception as e:
        logger.warning(f"Markdown→LaTeX 转换失败: {e}")
        return None

    # Step 2: xelatex 编译
    if not shutil.which("xelatex"):
        logger.warning("xelatex 未安装，跳过 PDF 编译")
        return None

    try:
        # 修复 MiKTeX 字体缓存
        subprocess.run(
            ["initexmf", "--mkmaps"], capture_output=True, timeout=30,
        )
        subprocess.run(
            ["initexmf", "--update-fndb"], capture_output=True, timeout=30,
        )
    except Exception:
        pass  # 非 MiKTeX 环境可忽略

    try:
        # 两次编译（不捕获输出，避免 MiKTeX 管道阻塞）
        for _ in range(2):
            subprocess.run(
                ["xelatex", "-interaction=nonstopmode", tex_path],
                timeout=180, cwd=wd,
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
            )
        if os.path.exists(pdf_path):
            logger.info(f"PDF 编译完成: {pdf_path}")
            return pdf_path
        return None
    except Exception as e:
        logger.warning(f"PDF 编译失败: {e}")
        return None


# 向后兼容
def latex_2_pdf(task_id: str, work_dir: str | None = None) -> str | None:
    """向后兼容：调用 md_2_pdf。"""
    return md_2_pdf(task_id, work_dir)


def _parse_html_tables(text: str) -> tuple[str, list[dict]]:
    """使用 BeautifulSoup 鲁棒解析 HTML 表格，返回 (清理后文本, 结构化表格数据)。

    相比正则方案，BeautifulSoup 能正确处理：
    - 嵌套标签（<sub>, <sup>, <b> 等）
    - 不闭合/不规范的 HTML
    - 多行表头、缺失单元格

    Args:
        text: 包含 HTML <table> 的 Markdown 文本。

    Returns:
        (cleaned_text, tables_data)
        - cleaned_text: HTML 表已替换为 Markdown 管道表的文本
        - tables_data: [{"caption": "表1: xxx", "headers": [...], "rows": [[...], ...]}, ...]
    """
    from bs4 import BeautifulSoup
    import re

    soup = BeautifulSoup(text, "lxml")
    tables_data = []

    for table in soup.find_all("table"):
        # ---- 提取表头 ----
        headers = []
        thead = table.find("thead")
        if thead:
            headers = [th.get_text(strip=True) for th in thead.find_all("th")]
        if not headers:
            first_row = table.find("tr")
            if first_row:
                # 第一行如果有 th 则视为表头，否则全 td 也当表头
                first_cells = first_row.find_all(["th", "td"])
                if first_cells:
                    headers = [c.get_text(strip=True) for c in first_cells]

        # ---- 提取数据行 ----
        all_rows = table.find_all("tr")
        data_start = 1 if headers and all_rows else 0
        rows = []
        for tr in all_rows[data_start:]:
            cells = tr.find_all(["td", "th"])
            if cells:
                rows.append([c.get_text(strip=True) for c in cells])

        if not rows:
            continue  # 空表跳过

        # 确保所有行列数一致
        max_cols = max(len(headers), max((len(r) for r in rows), default=0))
        if headers:
            while len(headers) < max_cols:
                headers.append("")
        for r in rows:
            while len(r) < max_cols:
                r.append("")

        # ---- 尝试匹配表名（<table> 前面的 **表X: xxx** 或 表X: xxx）----
        table_str = str(table)
        # 在原始文本中定位此 table 的位置
        pos = text.find(table_str)
        caption = ""
        if pos > 0:
            before = text[:pos]
            # 匹配前面的 **表X: xxx** 或 表X：xxx
            cap_match = re.search(
                r"(?:\*\*)?(?:表\d+[：:].*?)(?:\*\*)?\s*$", before, re.MULTILINE
            )
            if cap_match:
                caption = cap_match.group(0).strip().strip("*").strip()

        tables_data.append({
            "caption": caption,
            "headers": headers,
            "rows": rows,
        })

    # ---- 替换所有 HTML <table> 为 Markdown 管道表 ----
    def _html_to_md_table(table_el, td: dict):
        """将单个 BeautifulSoup table 元素 + 数据 渲染为 markdown 管道表。"""
        h = td["headers"]
        r = td["rows"]
        lines = []
        if h:
            lines.append("| " + " | ".join(h) + " |")
        lines.append("| " + " | ".join("---" for _ in range(max(len(h), len(r[0]) if r else 1))) + " |")
        for row in r:
            lines.append("| " + " | ".join(row) + " |")
        return "\n".join(lines)

    cleaned = text
    for table_el, td in zip(soup.find_all("table"), tables_data):
        md_table = _html_to_md_table(table_el, td)
        cleaned = cleaned.replace(str(table_el), "\n\n" + md_table + "\n\n", 1)

    return cleaned, tables_data


# ============================================================
# 方案一：HTML 表格直接用 python-docx 建 Word 表（单路径，无 fallback）
# ============================================================

def _extract_html_tables(md_path: str) -> list[dict]:
    """从 res.md 提取 HTML 表格，替换为占位符。

    BS4 解析 → 结构化数据 + 清洗后的 MD（pandoc 不再见到表格标签）。
    返回: [{"caption": "表1: xxx", "headers": [...], "rows": [[...], ...]}, ...]
    """
    import re
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    tables_data = []

    def _parse_and_store(match):
        html = match.group(0)
        # 找表号（前面最近的 **表X: xxx**）
        start_pos = match.start()
        before = content[max(0, start_pos - 200):start_pos]
        cap_match = re.search(r"\*\*表(\d+)[：:]([^*]+)\*\*", before)
        caption = f"表{cap_match.group(1)}: {cap_match.group(2).strip()}" if cap_match else ""

        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html, "lxml")
            table = soup.find("table")
            if not table:
                return "[TABLE_PLACEHOLDER]"

            headers = []
            rows = []
            for tr in table.find_all("tr"):
                cells = [td.get_text(strip=True) for td in tr.find_all(["th", "td"])]
                if not cells:
                    continue
                if not headers and tr.find("th"):
                    headers = cells
                else:
                    rows.append(cells)

            if not headers and rows:
                headers = rows[0]
                rows = rows[1:]

            tables_data.append({
                "caption": caption,
                "headers": headers,
                "rows": rows,
            })
            return f"\n[TABLE_{len(tables_data) - 1}]\n"
        except Exception:
            return "[TABLE_PLACEHOLDER]"

    content = re.sub(r"<table>.*?</table>", _parse_and_store, content, flags=re.DOTALL)

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(content)

    logger.info(f"HTML 表格提取: {len(tables_data)} 个 → 占位符，res.md 已清洗")
    return tables_data


def _extract_tables_from_json(work_dir: str) -> list[dict]:
    """从 res.json 恢复 HTML 表格数据（当 res.md 已被清洗时使用）。"""
    import json, re
    res_json = os.path.join(work_dir, "res.json")
    if not os.path.exists(res_json):
        return []

    with open(res_json, "r", encoding="utf-8") as f:
        res = json.load(f)

    tables = []
    # 按论文章节顺序遍历，确保 [TABLE_N] 索引匹配
    SEQ = ["symbol", "ques1", "ques2", "ques3", "ques4", "sensitivity_analysis", "judge"]
    for key in SEQ:
        if key not in res:
            continue
        content = str(res[key].get("response_content", ""))
        # 找所有 HTML 表格
        html_tables = re.findall(r"<table>(.*?)</table>", content, re.DOTALL)
        for html in html_tables:
            # 找表号（前面的 **表X: xxx**）
            idx = content.find(html) - 200
            before = content[max(0, idx):content.find(html)]
            cap_match = re.search(r"\*\*表(\d+)[：:]([^*]+)\*\*", before)
            caption = f"表{cap_match.group(1)}: {cap_match.group(2).strip()}" if cap_match else ""

            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(f"<table>{html}</table>", "lxml")
                table = soup.find("table")
                if not table:
                    continue
                headers, rows = [], []
                for tr in table.find_all("tr"):
                    cells = [td.get_text(strip=True) for td in tr.find_all(["th", "td"])]
                    if not cells:
                        continue
                    if not headers and tr.find("th"):
                        headers = cells
                    else:
                        rows.append(cells)
                if not headers and rows:
                    headers, rows = rows[0], rows[1:]
                tables.append({"caption": caption, "headers": headers, "rows": rows})
            except Exception:
                pass

    logger.info(f"从 res.json 恢复表格: {len(tables)} 个")
    return tables


def _latex_cell_to_text(text: str) -> str:
    """将表格单元格中的 LaTeX 数学公式转为 Word 可显示的 Unicode 文本。

    处理内容: $...$ 定界符 / Greek字母 / 上下标 / 常见数学符号
    """
    import re

    # 1. 移除 $ 定界符
    text = re.sub(r'\$(.+?)\$', r'\1', text)

    # 2. Greek 字母
    GREEK = {
        "\\alpha": "α", "\\beta": "β", "\\gamma": "γ", "\\delta": "δ",
        "\\epsilon": "ε", "\\zeta": "ζ", "\\eta": "η", "\\theta": "θ",
        "\\iota": "ι", "\\kappa": "κ", "\\lambda": "λ", "\\mu": "μ",
        "\\nu": "ν", "\\xi": "ξ", "\\pi": "π", "\\rho": "ρ",
        "\\sigma": "σ", "\\tau": "τ", "\\upsilon": "υ", "\\phi": "φ",
        "\\chi": "χ", "\\psi": "ψ", "\\omega": "ω",
        "\\Alpha": "Α", "\\Beta": "Β", "\\Gamma": "Γ", "\\Delta": "Δ",
        "\\Sigma": "Σ", "\\Omega": "Ω",
    }
    for latex, uni in GREEK.items():
        text = text.replace(latex, uni)

    # 3. 上标 ^X → ˣ
    SUPER = str.maketrans("0123456789+-=()abcdefghijklmnoprstuvwxyz",
                           "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ᵃᵇᶜᵈᵉᶠᵍʰⁱʲᵏˡᵐⁿᵒᵖʳˢᵗᵘᵛʷˣʸᶻ")
    def _replace_sup(m):
        return m.group(1).translate(SUPER)
    text = re.sub(r'\^\{?(-?\d+|\w)\}?', _replace_sup, text)

    # 4. 下标 _X → ₓ
    SUB = str.maketrans("0123456789+-=()abcdefghijklmnoprstuvwxyz",
                         "₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎ₐₑ𝒸𝒹ₑ𝒻𝓰ₕᵢⱼₖₗₘₙₒₚᵣₛₜᵤᵥ𝓌ₓᵧ𝓏")
    def _replace_sub(m):
        return m.group(1).translate(SUB)
    text = re.sub(r'_\{?(-?\d+|\w+)\}?', _replace_sub, text)

    # 5. 常见数学命令
    MATH_CMDS = {
        "\\infty": "∞", "\\pm": "±", "\\mp": "∓", "\\cdot": "·",
        "\\times": "×", "\\approx": "≈", "\\sim": "~", "\\leq": "≤",
        "\\geq": "≥", "\\neq": "≠", "\\equiv": "≡", "\\propto": "∝",
        "\\sum": "Σ", "\\prod": "Π", "\\int": "∫",
        "\\partial": "∂", "\\nabla": "∇", "\\forall": "∀", "\\exists": "∃",
        "\\in": "∈", "\\notin": "∉", "\\subset": "⊂", "\\subseteq": "⊆",
        "\\rightarrow": "→", "\\leftarrow": "←", "\\Rightarrow": "⇒",
        "\\bar": "", "\\hat": "", "\\tilde": "", "\\vec": "",
        "\\mathbf": "", "\\mathrm": "", "\\text": "", "\\textbf": "",
        "\\quad": " ", "\\qquad": "  ",
    }
    for latex, uni in MATH_CMDS.items():
        text = text.replace(latex, uni)

    # 6. 清理残留的大括号
    text = text.replace("{", "").replace("}", "")

    return text.strip()

def _insert_tables_into_docx(doc, tables_data: list[dict],
                              cn_font: str, cn_heading_font: str, en_font: str) -> int:
    """将 BS4 解析的表格数据插入 DOCX（替换占位符段落）。

    关键：不再使用 _set_cell_format（会破坏 MathML）。
    改为在空单元格上用 add_run 写入纯文本 + 设置字体。
    如果 Pandoc 已经在该位置生成了内容（正文段落），则保留不删。
    """
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not tables_data:
        return 0

    # 找到所有占位符段落
    import re
    placeholder_re = re.compile(r"\[TABLE_(\d+)\]")

    inserted = 0
    para_idx = 0
    while para_idx < len(doc.paragraphs):
        p = doc.paragraphs[para_idx]
        m = placeholder_re.match(p.text.strip())
        if not m:
            para_idx += 1
            continue

        table_idx = int(m.group(1))
        if table_idx >= len(tables_data):
            para_idx += 1
            continue

        td = tables_data[table_idx]
        headers = td.get("headers", [])
        rows = td.get("rows", [])
        if not headers and not rows:
            para_idx += 1
            continue

        all_rows = ([headers] if headers else []) + rows
        num_cols = max(len(r) for r in all_rows)
        for r in all_rows:
            while len(r) < num_cols:
                r.append("")

        # 创建 Word 表格
        ref_element = p._element
        parent = ref_element.getparent()
        idx_in_parent = list(parent).index(ref_element)

        tbl = OxmlElement("w:tbl")
        tblGrid = OxmlElement("w:tblGrid")
        col_w = int(9000 / max(num_cols, 1))
        for _ in range(num_cols):
            gc = OxmlElement("w:gridCol")
            gc.set(qn("w:w"), str(col_w))
            tblGrid.append(gc)
        tbl.append(tblGrid)

        # 三线表边框
        tblPr = OxmlElement("w:tblPr")
        tblBorders = OxmlElement("w:tblBorders")
        for bn, sz in [("top", 12), ("bottom", 12), ("left", 0), ("right", 0),
                        ("insideH", 0), ("insideV", 0)]:
            b = OxmlElement(f"w:{bn}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), str(sz))
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "000000")
            tblBorders.append(b)
        tblPr.append(tblBorders)
        tbl.append(tblPr)

        # 填充数据行
        for r_idx, row_data in enumerate(all_rows):
            tr = OxmlElement("w:tr")
            for c_idx, cell_text in enumerate(row_data):
                tc = OxmlElement("w:tc")
                # 表头下方横线
                if r_idx == 0:
                    tcPr = OxmlElement("w:tcPr")
                    tcBorders = OxmlElement("w:tcBorders")
                    bottom = OxmlElement("w:bottom")
                    bottom.set(qn("w:val"), "single")
                    bottom.set(qn("w:sz"), "6")
                    bottom.set(qn("w:space"), "0")
                    bottom.set(qn("w:color"), "000000")
                    tcBorders.append(bottom)
                    tcPr.append(tcBorders)
                    tc.append(tcPr)

                # 段落 + 文本
                wp = OxmlElement("w:p")
                r = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                # 字体
                rFonts = OxmlElement("w:rFonts")
                rFonts.set(qn("w:ascii"), en_font)
                rFonts.set(qn("w:eastAsia"), cn_heading_font if r_idx == 0 else cn_font)
                rPr.append(rFonts)
                # 字号
                sz_elem = OxmlElement("w:sz")
                sz_elem.set(qn("w:val"), "18")  # 9pt = 18 half-pt
                rPr.append(sz_elem)
                # 加粗（表头）
                if r_idx == 0:
                    b_elem = OxmlElement("w:b")
                    rPr.append(b_elem)
                r.append(rPr)
                t = OxmlElement("w:t")
                t.set(qn("xml:space"), "preserve")
                t.text = _latex_cell_to_text(cell_text)
                r.append(t)
                wp.append(r)
                tc.append(wp)
                tr.append(tc)
            tbl.append(tr)

        parent.insert(idx_in_parent, tbl)
        parent.remove(ref_element)
        inserted += 1
        # 不递增 para_idx（占位符已删除，下一个段落到当前位置）

    logger.info(f"Word 表格插入: {inserted} 个 (python-docx 直接建表)")
    return inserted


# ---- 旧版函数（保留供 res.md 使用）----

def _html_tables_to_markdown(md_path: str) -> None:
    """将 HTML <table> 转换为 markdown 管道表（使用 BeautifulSoup 鲁棒解析）。

    LLM 写 HTML 表格最可靠，但 pandoc 不解析 HTML <table>。
    此函数在 pandoc 前将 HTML 表转为 markdown 管道表。
    """
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    if "<table>" not in content and "<table " not in content:
        return  # 无 HTML 表格，无需处理

    cleaned, tables_data = _parse_html_tables(content)

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(cleaned)

    logger.info(f"HTML→Markdown 表格转换完成，共处理 {len(tables_data)} 个表格")


def _fix_markdown_tables(md_path: str) -> None:
    """修复 markdown 表格语法，确保 pandoc 能正确转换。

    将非标准分隔符（破折号、空白header等）标准化为 `| ---- |` 格式。
    """
    import re
    with open(md_path, "r", encoding="utf-8") as f:
        content = f.read()

    lines = content.split("\n")
    fixed = 0

    for i in range(1, len(lines) - 1):
        line = lines[i].strip()

        # 检测表格分隔行：以 | 开头，只含 | - : 空格
        if not (line.startswith("|") and line.endswith("|")):
            continue

        # 分隔行特征：去掉 | 后，剩余字符全是 - : 空格
        inner = line[1:-1]
        if not inner or not all(c in "-: | " for c in inner):
            continue

        # 上一行必须是表头行（以 | 开头）
        prev = lines[i - 1].strip()
        if not prev.startswith("|"):
            continue

        # 标准化分隔符：每个列用最少 4 个 `-`
        cols = [c.strip() for c in prev.split("|")[1:-1]]
        if not cols:
            continue

        new_sep = "| " + " | ".join("-" * max(4, len(c)) for c in cols) + " |"
        if lines[i].rstrip() != new_sep:
            lines[i] = new_sep
            fixed += 1

    if fixed > 0:
        with open(md_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        logger.info(f"表格语法修复: {fixed} 个分隔符已标准化")


def _sync_template_styles(template_path: str, target_path: str,
                          tables_data: list | None = None) -> None:
    """将专业学术格式注入生成的 DOCX。

    Pandoc 只创建段落不设样式，所有内容都是 Normal 样式的默认字体。
    此函数根据每段文本内容判断角色（标题/正文/图表标注），
    强制应用对应的字体、字号、加粗、行距、对齐方式。

    标准：中文宋体(正文)/黑体(标题)，英文 Times New Roman。

    Args:
        template_path: 参考模板 .docx 路径（保留，供未来使用）。
        target_path: 生成的 .docx 路径（原地修改）。
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.warning("python-docx 未安装，跳过格式注入")
        return

    try:
        doc = Document(target_path)
    except Exception as e:
        logger.warning(f"无法打开文档进行格式注入: {e}")
        return

    # 中文字体映射
    CN_FONT = "宋体"
    CN_HEADING_FONT = "黑体"
    EN_FONT = "Times New Roman"

    import re
    h1_pattern = re.compile(r"^[一二三四五六七八九十]、")     # "一、问题重述"
    h2_pattern = re.compile(r"^\d+\.\d+\s")                 # "1.1 背景"
    h3_pattern = re.compile(r"^\d+\.\d+\.\d+\s")            # "1.1.1 思路"
    abstract_pattern = re.compile(r"^摘要")                   # "摘要"
    fig_caption = re.compile(r"^(?:\*\*)?图\d+[：:）\)]")       # "**图1: xxx**" 或 "图1：xxx" (pandoc会去掉**)

    formatted = {"h1": 0, "h2": 0, "h3": 0, "body": 0, "title": 0, "abstract": 0}

    # 识别文档标题：第一个非空段落（pandoc 已去掉 # 前缀）
    title_idx = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() and not p.text.strip().startswith("!"):
            title_idx = i
            break

    # 表格数据已在 _extract_html_tables 中预处理（BS4 解析）

    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text:
            continue

        # 跳过图片占位符
        if text.startswith("!["):
            continue

        # 图注行：居中，宋体小五
        is_fig_caption = bool(fig_caption.match(text))

        # 判断段落角色
        is_title = (i == title_idx)
        is_abstract = abstract_pattern.match(text) and not is_title
        is_h1 = h1_pattern.match(text)
        is_h2 = h2_pattern.match(text)
        is_h3 = h3_pattern.match(text)

        if is_fig_caption:
            _apply_fig_caption_format(p, Pt(10), CN_FONT, EN_FONT)
            formatted["body"] += 1
        elif is_title:
            _apply_heading_format(p, Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, CN_HEADING_FONT, EN_FONT)
            formatted["title"] += 1
        elif is_abstract:
            _apply_heading_format(p, Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, CN_HEADING_FONT, EN_FONT)
            formatted["abstract"] += 1
        elif is_h1:
            _apply_heading_format(p, Pt(16), True, WD_ALIGN_PARAGRAPH.CENTER, CN_HEADING_FONT, EN_FONT)
            formatted["h1"] += 1
        elif is_h2:
            _apply_heading_format(p, Pt(14), True, None, CN_HEADING_FONT, EN_FONT)
            formatted["h2"] += 1
        elif is_h3:
            _apply_heading_format(p, Pt(12), True, None, CN_HEADING_FONT, EN_FONT)
            formatted["h3"] += 1
        else:
            _apply_body_format(p, Pt(12), CN_FONT, EN_FONT)
            formatted["body"] += 1

    # 正文段落首行缩进
    for para_idx, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if not text or text.startswith("!["):
            continue
        if para_idx == title_idx:
            continue
        if h1_pattern.match(text) or h2_pattern.match(text) or h3_pattern.match(text) or abstract_pattern.match(text):
            continue
        if p.paragraph_format.first_line_indent is None:
            p.paragraph_format.first_line_indent = Pt(24)

    # 优先：通过 doc.tables API 直接格式化 pandoc 生成的 Word 表格为三线表
    # 插入 Word 真表格（BS4 解析 HTML → python-docx 直接建表，单路径）
    table_count = _insert_tables_into_docx(
        doc, tables_data or [], CN_FONT, CN_HEADING_FONT, EN_FONT
    )

    doc.save(target_path)
    logger.info(f"格式注入完成: H1={formatted['h1']} H2={formatted['h2']} "
                f"H3={formatted['h3']} Body={formatted['body']} "
                f"Tables={table_count}")


def _format_docx_tables_as_sanxian(doc, cn_font, cn_heading_font, en_font):
    """通过 doc.tables API 直接将 pandoc 生成的 Word 表格格式化为三线表。

    相比正则解析段落文本的旧方案，此方法直接操作 Word 表格对象，
    不会受到 Markdown 语法变化或 pandoc 版本差异的影响。

    Args:
        doc: python-docx Document 对象。
        cn_font: 中文字体名（正文）。
        cn_heading_font: 中文字体名（表头）。
        en_font: 英文字体名。
    """
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not doc.tables:
        return

    formatted_count = 0
    for table in doc.tables:
        num_rows = len(table.rows)
        num_cols = len(table.columns)
        if num_rows < 1 or num_cols < 1:
            continue

        tbl = table._tbl

        # ---- 设置表格宽度为页面 100% ----
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")
        # 替换已有的 tblW
        existing_tblW = tblPr.find(qn("w:tblW"))
        if existing_tblW is not None:
            tblPr.remove(existing_tblW)
        tblPr.insert(0, tblW)

        # ---- 三线表边框 ----
        tblBorders = OxmlElement("w:tblBorders")
        for border_name, sz in [("top", 12), ("bottom", 12),
                                 ("left", 0), ("right", 0),
                                 ("insideH", 0), ("insideV", 0)]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), str(sz))
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            tblBorders.append(border)
        existing_borders = tblPr.find(qn("w:tblBorders"))
        if existing_borders is not None:
            tblPr.remove(existing_borders)
        tblPr.append(tblBorders)

        # ---- 格式化表头行（第一行） ----
        header_row = table.rows[0]
        for cell in header_row.cells:
            _set_cell_format(cell, cell.text, cn_heading_font, en_font, Pt(10), True, True)

        # ---- 格式化数据行 ----
        for row in table.rows[1:]:
            for cell in row.cells:
                _set_cell_format(cell, cell.text, cn_font, en_font, Pt(10), False, False)

        formatted_count += 1

    if formatted_count > 0:
        logger.info(f"三线表格式化: {formatted_count} 个 Word 表格已处理")


def _apply_fig_caption_format(p, size, cn_font, en_font):
    """图注格式：居中，宋体小五。"""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = size
        run.font.name = en_font
        run.font.bold = False
        _set_cn_font(run, cn_font)


def _apply_heading_format(p, size, bold, alignment, cn_font, en_font):
    """对段落应用标题格式。"""
    from docx.shared import Pt
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    if alignment is not None:
        p.alignment = alignment
    for run in p.runs:
        run.font.size = size
        run.font.bold = bold
        run.font.name = en_font
        _set_cn_font(run, cn_font)


def _apply_body_format(p, size, cn_font, en_font):
    """对段落应用正文字体格式。"""
    from docx.shared import Pt
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.size = size
        run.font.name = en_font
        _set_cn_font(run, cn_font)


def _find_markdown_tables(doc) -> list[tuple[int, int]]:
    """找到文档中所有表格段落范围。

    支持两种格式:
    1. Markdown 管道表: | col | col |
    2. Unicode 画线表: ┌────┬────┐（LLM 偶尔产出）

    返回 [(start_idx, end_idx), ...] 的列表。
    """
    import re
    pipe_row = re.compile(r"^\|.*\|$")
    pipe_sep = re.compile(r"^\|[\s\-:|]+\|$")
    # Unicode box-drawing characters
    box_chars = set("┌┬┐├┼┤└┴┘│─═")

    blocks = []
    i = 0
    paragraphs = doc.paragraphs
    while i < len(paragraphs):
        text = paragraphs[i].text.strip()

        # 格式1: Markdown 管道表
        if pipe_row.match(text):
            start = i
            i += 1
            while i < len(paragraphs):
                t = paragraphs[i].text.strip()
                if pipe_sep.match(t) or pipe_row.match(t):
                    i += 1
                else:
                    break
            if i > start + 1:
                blocks.append((start, i))
            continue

        # 格式2: Unicode 画线表
        is_box = any(c in box_chars for c in text)
        if is_box and len(text) > 10:
            start = i
            i += 1
            while i < len(paragraphs):
                t = paragraphs[i].text.strip()
                if any(c in box_chars for c in t) or (
                    "│" in t  # 数据行含有竖线分隔
                ):
                    i += 1
                else:
                    break
            if i > start + 1:
                blocks.append((start, i))
        else:
            i += 1

    return blocks


def _convert_markdown_tables(doc, table_blocks, cn_font, cn_heading_font, en_font):
    """将 markdown 表格段落转换为真正的 Word 三线表。"""
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if not table_blocks:
        return

    for start, end in reversed(table_blocks):  # 从后往前，避免索引偏移
        paragraphs = doc.paragraphs
        rows_data = []
        for i in range(start, end):
            text = paragraphs[i].text.strip()
            # 跳过分隔行（管道表分隔符 或 Unicode 画线行）
            if all(c in "|-: " for c in text):
                continue
            if any(c in "┌┬┐├┼┤└┴┘═" for c in text) and "│" not in text:
                continue  # Unicode 边框行
            # 解析单元格
            if "|" in text and text.startswith("|"):
                cells = [c.strip() for c in text.split("|")[1:-1]]
            elif "│" in text:
                cells = [c.strip() for c in text.split("│")[1:-1]]
            else:
                continue
            if cells and any(c for c in cells if c):  # 跳过空行
                rows_data.append(cells)

        if len(rows_data) < 2:
            continue

        # 在第一个表格段落位置插入 Word 表格
        ref_para = paragraphs[start]
        tbl = ref_para._element.getparent().makeelement(
            qn("w:tbl"), {}
        )
        ref_para._element.addprevious(tbl)

        # 创建表格
        from docx.table import Table
        table = Table(tbl, doc)

        # 计算列宽（均分）
        num_cols = len(rows_data[0])
        page_width = Cm(14.64)  # A4 减去边距
        col_width = page_width / num_cols

        header_row_data = rows_data[0]
        data_rows = rows_data[1:]

        # 设置表格属性
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)

        # 表格宽度
        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), str(int(page_width / 914400 * 5000)))  # 百分比
        tblW.set(qn("w:type"), "pct")
        tblPr.append(tblW)

        # 表格边框：三线表
        tblBorders = OxmlElement("w:tblBorders")
        for border_name, size in [("top", 12), ("bottom", 12), ("left", 0), ("right", 0),
                                   ("insideH", 0), ("insideV", 0)]:
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), str(size))
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "000000")
            tblBorders.append(border)
        tblPr.append(tblBorders)

        # 表头行
        if header_row_data:
            row = table.add_row()
            for j, cell_text in enumerate(header_row_data):
                cell = row.cells[j] if j < len(row.cells) else row.add_cell()
                _set_cell_format(cell, cell_text, cn_heading_font, en_font, Pt(10), True, True)

        # 数据行
        for row_data in data_rows:
            row = table.add_row()
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j] if j < len(row.cells) else row.add_cell()
                _set_cell_format(cell, cell_text, cn_font, en_font, Pt(10), False, False)

        # 删除原始 markdown 表格段落
        for i in range(start, end):
            p = paragraphs[i]
            p._element.getparent().remove(p._element)

    logger.info(f"表格转换: {len(table_blocks)} 个 markdown 表格 → Word 三线表")


def _convert_inline_tables(doc, cn_font, cn_heading_font, en_font):
    """处理 pandoc 未转换的表格——把含 | 分隔符的文本段落转为 Word 表。"""
    from docx.shared import Pt, Cm
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.table import Table
    import re

    converted = 0
    para_idx = 0
    while para_idx < len(doc.paragraphs):
        p = doc.paragraphs[para_idx]
        text = p.text.strip()

        # 检测 pandoc 转换失败的表格：段落内含多个 `|` 分隔
        if text.count("|") < 3 or not any(c.isdigit() or c.isalpha() for c in text if c != "|"):
            para_idx += 1
            continue

        # 尝试从文本中提取表格数据
        parts = text.split("|")
        # 找表头（前面的纯文本）和表数据
        caption = ""
        data_start = 0
        for i, part in enumerate(parts):
            part = part.strip()
            if re.match(r".*[:-–—]{3,}.*", part):  # 分隔符行
                data_start = i
                break

        if data_start == 0:
            para_idx += 1
            continue

        # 提取表头
        header_parts = [p.strip() for p in parts[:data_start] if p.strip()]
        if not header_parts:
            para_idx += 1
            continue

        # 表号从段落开头提取
        cap_match = re.match(r"^(?:\*\*)?(表\d+[：:][^*]+?)(?:\*\*)?\s*\|", text)
        if cap_match:
            caption = cap_match.group(1).strip()
            # 从第一个 | 后面开始解析
            table_text = text[cap_match.end()-1:]  # include the first |
        else:
            table_text = text

        # 解析表格行：用 | 分割
        table_parts = table_text.split("|")
        # 过滤掉表头行和分隔行之间的内容，找到数据行
        rows = []
        current_row = []
        for part in table_parts:
            part = part.strip()
            if re.match(r"^[\s\-–—:]+$", part):  # 分隔符行
                if current_row:
                    rows.append(current_row)
                    current_row = []
                continue
            if part:
                current_row.append(part)

        if current_row:
            rows.append(current_row)

        if len(rows) < 2:
            para_idx += 1
            continue

        # 创建 Word 表格（用 doc.add_table 标准方式）
        num_cols = max(len(r) for r in rows)
        for r in rows:
            while len(r) < num_cols:
                r.append("")

        # 在段落前插入表格
        ref_element = p._element
        parent = ref_element.getparent()
        tbl_element = OxmlElement("w:tbl")

        # 添加 tblGrid（必需）
        tblGrid = OxmlElement("w:tblGrid")
        col_width = int(9000 / num_cols)  # DXA units
        for _ in range(num_cols):
            gridCol = OxmlElement("w:gridCol")
            gridCol.set(qn("w:w"), str(col_width))
            tblGrid.append(gridCol)
        tbl_element.append(tblGrid)

        parent.insert(list(parent).index(ref_element), tbl_element)
        table = Table(tbl_element, doc)

        # 三线表边框
        tblPr = tbl_element.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl_element.insert(0, tblPr)

        tblW = OxmlElement("w:tblW")
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")
        tblPr.append(tblW)

        tblBorders = OxmlElement("w:tblBorders")
        for bn, sz in [("top", 12), ("bottom", 12), ("left", 0), ("right", 0),
                        ("insideH", 0), ("insideV", 0)]:
            b = OxmlElement(f"w:{bn}")
            b.set(qn("w:val"), "single")
            b.set(qn("w:sz"), str(sz))
            b.set(qn("w:space"), "0")
            b.set(qn("w:color"), "000000")
            tblBorders.append(b)
        tblPr.append(tblBorders)

        # 填充表格
        for r_idx, row_data in enumerate(rows):
            row = table.add_row()
            for c_idx, cell_text in enumerate(row_data):
                if c_idx >= len(row.cells):
                    row.add_cell()
                cell = row.cells[c_idx]
                is_header = (r_idx == 0)
                _set_cell_format(
                    cell, cell_text,
                    cn_heading_font if is_header else cn_font,
                    en_font, Pt(9), is_header, is_header
                )

        # 删除原始文本段落
        p._element.getparent().remove(p._element)
        converted += 1
        # 不递增 para_idx，因为当前索引已被删除

    if converted:
        logger.info(f"内联表格转换: {converted} 个文本表格 → Word 表")


def _set_cell_format(cell, text, cn_font, en_font, size, bold, is_header):
    """设置单元格内容和格式。"""
    from docx.shared import Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # 清空单元格
    for p in cell.paragraphs:
        p._element.getparent().remove(p._element)

    # 添加段落
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

    run = p.add_run(text)
    run.font.size = size
    run.font.name = en_font
    run.font.bold = bold
    _set_cn_font(run, cn_font)

    # 单元格边框：只有表头下方有横线
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    if is_header:
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "0")
        bottom.set(qn("w:color"), "000000")
        tcBorders.append(bottom)
    tcPr.append(tcBorders)


def _set_cn_font(run, cn_font):
    """设置 run 的中文字体（通过 XML）。"""
    rpr = run._r.get_or_add_rPr()
    from lxml import etree
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    rFonts = rpr.find(".//w:rFonts", nsmap)
    if rFonts is None:
        rFonts = etree.SubElement(rpr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
    rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", cn_font)


def split_footnotes(text: str) -> tuple[str, list[tuple[str, str]]]:
    """从文本中分离正文和脚注。

    Args:
        text: 包含脚注的完整文本。

    Returns:
        (正文, 脚注列表) 的元组，脚注格式为 (编号, 内容)。
    """
    main_text = re.sub(
        r"\n\[\^\d+\]:.*?(?=\n\[\^|\n\n|\Z)", "", text, flags=re.DOTALL
    ).strip()

    # 匹配脚注定义
    footnotes = re.findall(r"\[\^(\d+)\]:\s*(.+?)(?=\n\[\^|\n\n|\Z)", text, re.DOTALL)
    logger.info(f"main_text:{main_text} \n footnotes:{footnotes}")
    return main_text, footnotes
